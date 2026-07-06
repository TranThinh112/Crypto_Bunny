from __future__ import annotations

import json
import sqlite3
from datetime import datetime, timedelta, timezone
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
CONFIG_PATH = ROOT / "config.local.telegram.yaml"
DB_PATH = ROOT / "data" / "bot_state.sqlite"
OUTPUT_PATH = ROOT / "BaoCao" / "TongQuatMini.docx"
LOCAL_TZ = timezone(timedelta(hours=7))


def local_label(value: str | None) -> str:
    if not value:
        return "-"
    try:
        dt = datetime.fromisoformat(str(value).replace("Z", "+00:00"))
        if dt.tzinfo is None:
            dt = dt.replace(tzinfo=timezone.utc)
        return dt.astimezone(LOCAL_TZ).strftime("%d/%m/%Y %H:%M:%S")
    except Exception:
        return str(value)


def read_nested_yaml_value(lines: list[str], path: list[str]) -> str | None:
    indent_step = 2
    start = 0
    for depth, key in enumerate(path[:-1]):
        target_indent = depth * indent_step
        found = False
        for idx in range(start, len(lines)):
            line = lines[idx]
            if not line.strip() or line.lstrip().startswith("#"):
                continue
            indent = len(line) - len(line.lstrip(" "))
            if indent == target_indent and line.strip() == f"{key}:":
                start = idx + 1
                found = True
                break
        if not found:
            return None

    final_key = path[-1]
    target_indent = (len(path) - 1) * indent_step
    for idx in range(start, len(lines)):
        line = lines[idx]
        if not line.strip() or line.lstrip().startswith("#"):
            continue
        indent = len(line) - len(line.lstrip(" "))
        stripped = line.strip()
        if indent < target_indent:
            break
        if indent == target_indent and stripped.startswith(f"{final_key}:"):
            return stripped.split(":", 1)[1].strip()
    return None


def load_config_snapshot() -> dict[str, str]:
    lines = CONFIG_PATH.read_text(encoding="utf-8").splitlines()
    keys = {
        "mode": ["mode"],
        "market_scan_source_symbols": ["ai", "internal", "market_scan_source_symbols"],
        "market_scan_max_symbols": ["ai", "internal", "market_scan_max_symbols"],
        "market_scan_min_win_probability_pct": ["ai", "internal", "market_scan_min_win_probability_pct"],
        "market_scan_pending_limit": ["ai", "internal", "market_scan_pending_limit"],
        "market_scan_use_ai": ["ai", "internal", "market_scan_use_ai"],
        "market_scan_require_ai_for_pending": ["ai", "internal", "market_scan_require_ai_for_pending"],
        "lc_pipeline_min_win_probability_pct": ["ai", "internal", "lc_pipeline_min_win_probability_pct"],
        "pending_review_min_confidence": ["pending_orders", "review", "min_confidence"],
        "pending_review_min_win_probability_pct": ["pending_orders", "review", "min_win_probability_pct"],
        "pending_review_min_risk_reward": ["pending_orders", "review", "min_risk_reward"],
        "pending_review_max_entry_drift_pct": ["pending_orders", "review", "max_entry_drift_pct"],
    }
    output: dict[str, str] = {}
    for name, path in keys.items():
        output[name] = read_nested_yaml_value(lines, path) or "-"
    return output


def load_state() -> tuple[dict[str, Any], list[str]]:
    con = sqlite3.connect(DB_PATH)
    con.row_factory = sqlite3.Row
    try:
        row = con.execute(
            "SELECT value FROM journal_state WHERE key = 'lc_internal_pipeline_state'"
        ).fetchone()
        state = json.loads(row["value"]) if row and row["value"] else {}
        keys = [item["key"] for item in con.execute("SELECT key FROM journal_state ORDER BY key").fetchall()]
        return state, keys
    finally:
        con.close()


def add_bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def main() -> None:
    config = load_config_snapshot()
    state, journal_keys = load_state()
    latest_scan = state.get("latest_mini_scan") or {}
    internal_lc = state.get("internal_lc") or []
    undecided = state.get("undecided") or []
    internal_notifications = state.get("internal_notifications") or []

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.add_run("TỔNG QUÁT PHẦN SCAN, LỌC VÀ QUYẾT ĐỊNH CỦA MINI")
    run.bold = True
    run.font.size = Pt(15)

    stamp = doc.add_paragraph()
    stamp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    stamp.add_run(
        "Dữ liệu tổng hợp tại thời điểm tạo báo cáo: "
        + datetime.now(LOCAL_TZ).strftime("%d/%m/%Y %H:%M:%S")
    )

    doc.add_heading("1. Tóm tắt nhanh", level=1)
    doc.add_paragraph(
        "Luồng hiện tại đã được đơn giản hóa theo đúng rule mới: chỉ giữ các cặp có win rate từ 62% trở lên, "
        "sau đó sắp xếp giảm dần theo win rate để xác định mức độ ưu tiên. Mini không còn chờ đủ 3 cặp mới chạy; "
        "chỉ cần pool có từ 1 đến 3 cặp hợp lệ là mini có thể duyệt."
    )
    doc.add_paragraph(
        f"Tại thời điểm kiểm tra gần nhất, mini scan có trạng thái '{latest_scan.get('status', '-')}', "
        f"tạo lúc {local_label(latest_scan.get('created_at'))}, "
        f"pool hiện tại = {latest_scan.get('pool_symbols') or []}, "
        f"selected = {latest_scan.get('selected_symbols') or []}. "
        f"LC nội bộ hiện có {len(internal_lc)} cặp, Chưa duyệt hiện có {len(undecided)} cặp."
    )

    doc.add_heading("2. Flow scan hiện tại", level=1)
    add_bullet(
        doc,
        f"Scan gốc lấy tối đa {config['market_scan_source_symbols']} cặp từ universe hiện hành để build candidate."
    )
    add_bullet(
        doc,
        "Mỗi candidate được chấm từ nhiều lớp: EMA, RSI, vị trí giá trong range, volume ratio, khung lớn, mẫu nến, tin tức và Market Guard."
    )
    add_bullet(
        doc,
        f"Lọc 1h: chỉ giữ candidate có win rate >= {config['lc_pipeline_min_win_probability_pct']}%, "
        "loại các cặp bị blocked do đã có vị thế/lệnh mở, rồi sort theo win rate giảm dần."
    )
    add_bullet(
        doc,
        "Nếu win rate bằng nhau thì tiếp tục so confidence, rồi volume ratio để phá hòa."
    )
    add_bullet(
        doc,
        "Tổng hợp 2h: gộp 2 cửa sổ 1h gần nhất, loại trùng symbol, lấy tối đa 3 cặp mạnh nhất làm LC nội bộ."
    )
    add_bullet(
        doc,
        "Nếu một symbol đã có vị thế hoặc đã có pending/active phù hợp thì bị loại khỏi pipeline trước khi giữ LC."
    )

    doc.add_heading("3. LC nội bộ và pool gửi lên mini", level=1)
    add_bullet(
        doc,
        "LC nội bộ là danh sách cặp sống sau bước tổng hợp 2h, tối đa 3 cặp."
    )
    add_bullet(
        doc,
        "Pool mini 4h lấy trực tiếp từ LC nội bộ hiện tại, không lấy lung tung từ toàn bộ candidate bên ngoài pool."
    )
    add_bullet(
        doc,
        "Rule mới: pool có 1, 2 hoặc 3 cặp đều hợp lệ; chỉ khi pool = 0 thì mini mới ở trạng thái waiting_lc."
    )
    add_bullet(
        doc,
        "Thứ tự ưu tiên trong pool bây giờ hoàn toàn bám theo win rate giảm dần của các cặp đã qua ngưỡng 62%."
    )

    doc.add_heading("4. Điều kiện để mini được duyệt", level=1)
    add_bullet(doc, "Có ít nhất 1 cặp trong pool LC nội bộ.")
    add_bullet(doc, "Mini scan phải tạo được selected_symbols, tức là có cặp được chọn hợp lệ.")
    add_bullet(
        doc,
        f"Nếu đang bật AI ({config['market_scan_use_ai']}), mini phải trả review hợp lệ; "
        f"nếu đang bắt buộc AI trước khi tạo pending ({config['market_scan_require_ai_for_pending']}), "
        "thì không được rơi vào fallback hoặc lỗi ai_review."
    )
    add_bullet(
        doc,
        f"Mini hiện chỉ được chọn tối đa {config['market_scan_pending_limit']} cặp trong một lượt, nên thực tế hệ thống chỉ đi tiếp 1 cặp mỗi vòng."
    )

    doc.add_heading("5. Điều kiện để cặp mini chọn đi tiếp lên LC hoặc LC_OKX", level=1)
    add_bullet(
        doc,
        "Phải map lại được candidate thật từ setup đã lưu trong LC nội bộ hoặc từ candidate hiện hành."
    )
    add_bullet(
        doc,
        "Không được trùng với cặp đang pending hoặc đang active."
    )
    add_bullet(
        doc,
        f"Phải qua review trước khi tạo pending: confidence >= {config['pending_review_min_confidence']}, "
        f"win rate >= {config['pending_review_min_win_probability_pct']}%, "
        f"RR >= {config['pending_review_min_risk_reward']}."
    )
    add_bullet(
        doc,
        "Ngoài ra còn bị chặn nếu spread quá lớn, stop loss quá ngắn/quá xa, lệch entry quá mức, cooldown chưa hết, chạm giới hạn lệnh ngày, vượt rủi ro ngày hoặc không xác minh được trạng thái OKX."
    )
    add_bullet(
        doc,
        "Nếu mode không phải dry_run/demo nội bộ và bước submit OKX thành công thì pending được ghi nhận với trạng thái LC_OKX."
    )

    doc.add_heading("6. Khi nào mini không gửi được lệnh lên OKX", level=1)
    add_bullet(doc, "Pool bằng 0 nên mini không chạy.")
    add_bullet(doc, "Mini không chọn ra selected_symbols.")
    add_bullet(doc, "AI review lỗi, fallback hoặc chưa có ai_review trong khi rule đang bắt buộc AI.")
    add_bullet(doc, "Candidate không map lại được từ setup đã lưu.")
    add_bullet(doc, "Candidate bị loại ở bước review rủi ro.")
    add_bullet(doc, "Quantity không hợp lệ hoặc submit OKX lỗi.")

    doc.add_heading("7. Vì sao trước đây nhìn thấy dữ liệu bị lệch", level=1)
    doc.add_paragraph(
        "Phần gây nhầm nhiều nhất là khác biệt giữa dữ liệu quyết định hiện hành và lịch sử thông báo dạng text. "
        "LC nội bộ, undecided, hourly_windows, two_hour_windows là state sống để ra quyết định. "
        "Còn internal_notifications là snapshot text lịch sử đã gửi. Vì vậy có thể thấy thông báo cũ vẫn còn tên cặp, "
        "trong khi state sống hiện tại đã rỗng hoặc đã bị lọc lại theo rule mới."
    )

    doc.add_heading("8. Trạng thái hệ thống lúc tạo báo cáo", level=1)
    add_bullet(doc, f"Mode hiện tại: {config['mode']}.")
    add_bullet(doc, f"Ngưỡng mini/local scan hiện hành: >= {config['market_scan_min_win_probability_pct']}%.")
    add_bullet(doc, f"Ngưỡng giữ LC nội bộ hiện hành: >= {config['lc_pipeline_min_win_probability_pct']}%.")
    add_bullet(doc, f"Mini chọn tối đa mỗi lượt: {config['market_scan_pending_limit']} cặp.")
    add_bullet(doc, f"Mini scan gần nhất: {latest_scan.get('status', '-')} lúc {local_label(latest_scan.get('created_at'))}.")
    add_bullet(doc, f"Pool gần nhất: {latest_scan.get('pool_symbols') or []}.")
    add_bullet(doc, f"Selected gần nhất: {latest_scan.get('selected_symbols') or []}.")
    add_bullet(doc, f"Số LC nội bộ hiện tại: {len(internal_lc)}.")
    add_bullet(doc, f"Số Chưa duyệt hiện tại: {len(undecided)}.")
    add_bullet(doc, f"Số thông báo nội bộ đã lưu: {len(internal_notifications)}.")
    add_bullet(
        doc,
        f"Deprecated key ai_internal_market_scan_latest còn tồn tại: {'có' if 'ai_internal_market_scan_latest' in journal_keys else 'không'}."
    )

    doc.add_heading("9. Kết luận ngắn", level=1)
    doc.add_paragraph(
        "Flow hiện tại đã rõ hơn trước: scan tạo candidate, 1h giữ các cặp >=62%, 2h tổng hợp thành LC nội bộ, "
        "pool mini lấy 1-3 cặp từ LC nội bộ, mini chọn tối đa 1 cặp, sau đó qua review rồi mới tạo LC hoặc LC_OKX. "
        "Điểm quan trọng nhất của bản cập nhật này là hệ thống đã bỏ hẳn rule ưu tiên 80%, thay bằng một logic đơn giản và nhất quán hơn: "
        "các cặp đủ ngưỡng 62% sẽ được xếp theo win rate giảm dần để xác định ưu tiên thật."
    )

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)


if __name__ == "__main__":
    main()
